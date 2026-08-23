from io import BytesIO

from docx import Document

from qpc.docx_exporter import render_docx
from qpc.schemas import (
    GeneratedPaper,
    GeneratedQuestion,
    GeneratedSection,
    PaperBlueprint,
    PaperMetadata,
    QuestionType,
    SectionBlueprint,
)


def _document_text(document: Document) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraphs + table_cells)


def test_render_docx_uses_calculated_total_in_header():
    blueprint = PaperBlueprint(
        metadata=PaperMetadata(
            grade="VII",
            subject="Social Science",
            exam_name="First Periodic Assessment 2026-27",
            date="17.07.2026",
            duration="2 Hour",
        ),
        sections=[
            SectionBlueprint(
                label="Section A",
                heading="Multiple Choice Based Questions",
                question_type=QuestionType.MCQ,
                questions_to_generate=1,
                questions_to_answer=1,
                marks_per_question=7,
                instruction="Choose the correct option.",
            )
        ],
    )
    paper = GeneratedPaper(
        sections=[
            GeneratedSection(
                label="Section A",
                heading="Multiple Choice Based Questions",
                question_type=QuestionType.MCQ,
                questions=[
                    GeneratedQuestion(
                        question_type=QuestionType.MCQ,
                        text="Which physical feature is known as the storehouse of water?",
                        options=["Coastal Plains", "Himalayas", "Deccan Plateau", "Islands"],
                    )
                ],
            )
        ]
    )

    data = render_docx(blueprint, paper)
    document = Document(BytesIO(data))
    text = _document_text(document)
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]

    assert "ZENITH PUBLIC SCHOOL" in text
    assert "First Periodic Assessment 2026-27" in text
    assert "Grade: VII" in text
    assert "Subject: - Social Science" in text
    assert "Date: - 17.07.2026" in text
    assert "M.M. 7        Time- 2 Hour" in text
    assert "Section A" in text
    assert "Choose the correct option." in text
    assert "Which physical feature" in text
    assert "a) Coastal Plains" in paragraph_texts
    assert "b) Himalayas" in paragraph_texts
    assert "c) Deccan Plateau" in paragraph_texts
    assert "d) Islands" in paragraph_texts
    assert "*****" in text


def test_render_docx_includes_case_study_instruction_passage_and_subquestions():
    blueprint = PaperBlueprint(
        metadata=PaperMetadata(
            grade="VII",
            subject="Social Science",
            exam_name="First Periodic Assessment 2026-27",
            date="17.07.2026",
            duration="2 Hour",
        ),
        sections=[
            SectionBlueprint(
                label="Section E",
                heading="Case Study Based Questions",
                question_type=QuestionType.CASE_STUDY,
                questions_to_generate=1,
                questions_to_answer=1,
                marks_per_question=4,
                instruction="Read the passage and answer the questions that follow.",
            )
        ],
    )
    paper = GeneratedPaper(
        sections=[
            GeneratedSection(
                label="Section E",
                heading="Case Study Based Questions",
                question_type=QuestionType.CASE_STUDY,
                passage="The northern mountains trap moisture-laden winds and feed major rivers.",
                questions=[
                    GeneratedQuestion(
                        question_type=QuestionType.CASE_STUDY,
                        text="Study the passage carefully.",
                        sub_questions=[
                            "Name the mountain range described.",
                            "State one benefit of these mountains.",
                        ],
                    )
                ],
            )
        ]
    )

    data = render_docx(blueprint, paper)
    document = Document(BytesIO(data))
    text = _document_text(document)

    assert "Section E" in text
    assert "Read the passage and answer the questions that follow." in text
    assert "The northern mountains trap moisture-laden winds" in text
    assert "1. Study the passage carefully." in text
    assert "(1) Name the mountain range described." in text
    assert "(2) State one benefit of these mountains." in text
