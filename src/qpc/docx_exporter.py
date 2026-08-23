from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from qpc.schemas import (
    GeneratedPaper,
    GeneratedSection,
    PaperBlueprint,
    QuestionType,
    SectionBlueprint,
)


def render_docx(blueprint: PaperBlueprint, paper: GeneratedPaper) -> bytes:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    _add_header(document, blueprint)
    sections_by_label = {section.label: section for section in paper.sections}
    for section_blueprint in blueprint.sections:
        generated_section = sections_by_label[section_blueprint.label]
        _add_section(document, section_blueprint, generated_section)

    end = document.add_paragraph("*****")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _add_header(document: Document, blueprint: PaperBlueprint) -> None:
    metadata = blueprint.metadata
    school = document.add_paragraph(metadata.school_name)
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    school.runs[0].bold = True
    school.runs[0].font.size = Pt(14)

    address = document.add_paragraph(metadata.school_address)
    address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation = document.add_paragraph(f"({metadata.affiliation})")
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER

    exam = document.add_paragraph(metadata.exam_name)
    exam.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exam.runs[0].bold = True

    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = f"Grade: {metadata.grade}"
    table.cell(0, 1).text = f"Subject: - {metadata.subject}"
    table.cell(1, 0).text = "Name: -_____________________________"
    table.cell(1, 1).text = f"Date: - {metadata.date}"
    table.cell(2, 0).text = "Roll No: - _______"
    table.cell(2, 1).text = (
        f"M.M. {blueprint.total_marks()}        Time- {metadata.duration}"
    )


def _section_summary(section: SectionBlueprint) -> str:
    total = section.section_marks()
    if section.questions_to_answer < section.questions_to_generate:
        return (
            f"{section.heading} (Any {section.questions_to_answer}) "
            f"({section.questions_to_answer}x{section.marks_per_question}={total})"
        )
    return (
        f"{section.heading} "
        f"({section.questions_to_answer}x{section.marks_per_question}={total})"
    )


def _add_section(
    document: Document,
    section_blueprint: SectionBlueprint,
    generated_section: GeneratedSection,
) -> None:
    label = document.add_paragraph(section_blueprint.label)
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.runs[0].bold = True

    summary = document.add_paragraph(_section_summary(section_blueprint))
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.runs[0].bold = True

    if section_blueprint.instruction.strip():
        document.add_paragraph(section_blueprint.instruction.strip())

    if generated_section.passage.strip():
        document.add_paragraph(generated_section.passage.strip())

    for index, question in enumerate(generated_section.questions, start=1):
        document.add_paragraph(f"{index}. {question.text.strip()}")
        if section_blueprint.question_type is QuestionType.MCQ:
            for option_index, option in enumerate(question.options):
                document.add_paragraph(f"{chr(97 + option_index)}) {option}")
        elif section_blueprint.question_type is QuestionType.MATCH:
            for left, right in question.pairs:
                document.add_paragraph(f"{left} - {right}")
        for sub_index, sub_question in enumerate(question.sub_questions, start=1):
            document.add_paragraph(f"({sub_index}) {sub_question}")
